from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# --- Helper: remove space before/after paragraph ---
def tight(para, space_before=0, space_after=0):
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)

# --- Helper: add a horizontal rule ---
def add_hr(doc):
    p = doc.add_paragraph()
    tight(p, 2, 2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

# --- Helper: bold run ---
def bold_run(para, text, size=None, color=None):
    run = para.add_run(text)
    run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

# --- Helper: normal run ---
def normal_run(para, text, size=10, italic=False, color=None):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

# ============================================================
# NAME
# ============================================================
name_para = doc.add_paragraph()
name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(name_para, 0, 2)
r = name_para.add_run('JP Hebert')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x1F, 0x1F, 0x1F)

# Contact line
contact_para = doc.add_paragraph()
contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(contact_para, 0, 2)
normal_run(contact_para, 'jphebert32@gmail.com  |  303-880-4657  |  Littleton, CO  |  linkedin.com/in/jp-hebert', size=9.5)

# Tagline
tag_para = doc.add_paragraph()
tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tight(tag_para, 2, 4)
r = tag_para.add_run('Senior BI Analyst  |  Power BI  |  Data Visualization  |  SQL  |  Snowflake')
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

add_hr(doc)

# ============================================================
# PROFESSIONAL SUMMARY
# ============================================================
def section_heading(doc, text):
    p = doc.add_paragraph()
    tight(p, 6, 2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

section_heading(doc, 'Professional Summary')

summary = doc.add_paragraph()
tight(summary, 0, 4)
normal_run(summary, (
    'Senior BI Analyst with 10+ years of experience and a proven track record of owning analytics end-to-end '
    'for high-scale products. Served as the sole BI analyst for T-Life — T-Mobile’s super app with 75M+ downloads '
    '— designing and maintaining 30 executive-facing Power BI dashboards accessed daily by VP of Product, Senior '
    'Directors, and C-suite stakeholders. Expert in the full pipeline: raw Snowflake/Databricks data to polished, '
    'decision-ready visuals. Seeking a Senior BI or BI Lead role where I can own the reporting layer and deliver '
    'dashboards that drive real product and business decisions.'
), size=10)

add_hr(doc)

# ============================================================
# CORE COMPETENCIES
# ============================================================
section_heading(doc, 'Core Competencies')

comp_para = doc.add_paragraph()
tight(comp_para, 0, 4)
normal_run(comp_para, (
    'Power BI (5+ years, Expert) ·  DAX ·  SQL (Advanced) ·  Snowflake ·  '
    'Databricks ·  Azure ·  Python ·  dbt ·  Metabase ·  '
    'KPI Framework Design ·  HEART Metrics ·  Executive Dashboard Delivery ·  '
    'Data Modeling ·  Stakeholder Management ·  Sprinklr ·  '
    'Adobe Creative Suite (InDesign, Illustrator, Photoshop)'
), size=10)

add_hr(doc)

# ============================================================
# PROFESSIONAL EXPERIENCE
# ============================================================
section_heading(doc, 'Professional Experience')

def job_header(doc, title, dates, company_line):
    p = doc.add_paragraph()
    tight(p, 4, 0)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10.5)
    # Right-align dates using tab
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    r2 = p.add_run('\t' + dates)
    r2.font.size = Pt(10)
    r2.bold = False

    p2 = doc.add_paragraph()
    tight(p2, 0, 2)
    r3 = p2.add_run(company_line)
    r3.italic = True
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def bullet(doc, text, indent=0.2):
    p = doc.add_paragraph(style='List Bullet')
    tight(p, 0, 1)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p

# --- Job 1: T-Sentiment ---
job_header(doc,
    'Senior Data Analyst',
    'Aug 2025 – Apr 2026',
    'ProCogia  |  Client: T-Mobile — Office of the President / VP Marketing  |  Remote')

bullet(doc, 'Sole BI analyst for T-Sentiment, T-Mobile’s executive social listening platform — tracked brand perception and customer sentiment across 7 platforms (Facebook, X, Reddit, Bluesky, Glassdoor, Instagram, YouTube, News) for the VP of Marketing and Office of the President.')
bullet(doc, 'Built and maintained Power BI dashboards translating social media data into daily executive reports: post categorization, platform benchmarks, custom alerts, and trend summaries delivered directly to C-suite.')
bullet(doc, 'Wrote complex SQL against Snowflake and Databricks to ingest, transform, and model large-scale social datasets into reliable reporting pipelines.')
bullet(doc, 'Developed Python scripts to automate data collection, sentiment scoring, and report generation — enabling near real-time delivery of insights to senior leadership.')
bullet(doc, 'Designed executive presentation decks and one-pagers synthesizing competitor benchmarking, sentiment trends, and engagement KPIs for senior marketing and strategy stakeholders.')

# --- Job 2: T-Life ---
job_header(doc,
    'Senior BI Analyst',
    'Sep 2021 – Aug 2025',
    'ProCogia  |  Client: T-Mobile — Product (VP of Product / Sr. Director of Product)  |  Remote')

bullet(doc, 'Served as the sole BI analyst for T-Life, T-Mobile’s super app (75M+ downloads, 4.79★ App Store rating) — owned the full analytics pipeline from raw Snowflake/Databricks data to polished executive dashboards.')
bullet(doc, 'Built and maintained 30 Power BI dashboards covering the HEART framework (Happiness, Engagement, Adoption, Retention, Task Success, Release Insights) and 5 product areas (Billing/Payments, Device Upgrades, SyncUP, Device Details/Add-Ons, Benefits) — accessed daily by 100+ stakeholders including VP of Product, Senior Directors, and Product Managers.')
bullet(doc, 'Presented monthly release dashboards and product performance reviews to C-suite; served as the source of truth for all T-Life product owners and release decisions.')
bullet(doc, 'Designed KPI frameworks and data models from scratch for greenfield product areas — built complete reporting infrastructure where none had previously existed.')
bullet(doc, 'Wrote advanced SQL against Snowflake and Databricks to extract, transform, and model product usage data supporting both self-serve analytics and executive reporting pipelines.')
bullet(doc, 'Prior to T-Life: owned BI reporting for SyncUP Products (Drive, Kids Watch, Tracker, Pets) for the Sr. Director of Product — subscriber retention, feature adoption funnels, and engagement KPIs across T-Mobile’s IoT suite.')

# --- Job 3: Network Tech ---
job_header(doc,
    'Data Visualization Specialist',
    'Dec 2020 – Sep 2021',
    'ProCogia  |  Client: T-Mobile — Network Technology  |  Remote')

bullet(doc, 'Designed and delivered executive-facing dashboards, VP-level presentation decks, infographics, and one-pagers for T-Mobile’s Network Technology organization — communicating complex network KPIs and performance trends to non-technical leadership.')
bullet(doc, 'Produced polished data visualizations and branded reports using Power BI, Tableau, and Adobe Creative Suite (InDesign, Illustrator, Photoshop).')
bullet(doc, 'Queried and analyzed large network datasets using SQL and Python to identify performance trends and surface anomalies for data-driven decision-making.')

# --- Job 4: MWD ---
job_header(doc,
    'High-Frequency Trading Data Analyst',
    'Jan 2011 – Dec 2019',
    'MWD Trading Inc.  |  Littleton, CO')

bullet(doc, 'Operated and continuously optimized commodities trading algorithms using SQL and Python across CME, ICE, and Eris markets, responding to dynamic market conditions in real time.')
bullet(doc, 'Built and maintained reporting and monitoring systems tracking P&L and performance for 3 trading groups, 12 traders, and 35 strategies — presented weekly to C-level executives.')
bullet(doc, 'Expanded the active trading strategy count from 4 to 20+.')
bullet(doc, 'Trained junior trade operators on algorithm behavior, risk management, and compliance with global market regulations.')

add_hr(doc)

# ============================================================
# EDUCATION & TRAINING
# ============================================================
section_heading(doc, 'Education & Training')

p = doc.add_paragraph()
tight(p, 2, 0)
bold_run(p, 'B.S. International Marketing', size=10)
normal_run(p, '  |  Johnson & Wales University  |  Denver, CO', size=10)

p2 = doc.add_paragraph()
tight(p2, 4, 0)
bold_run(p2, 'Data Analytics Intensive Bootcamp', size=10)
normal_run(p2, '  |  University of Denver  |  Sep 2020', size=10)

p3 = doc.add_paragraph()
tight(p3, 0, 0)
normal_run(p3, '24-week program: SQL, Python, Tableau, Machine Learning, Statistics, JavaScript, data engineering fundamentals.', size=10)

p4 = doc.add_paragraph()
tight(p4, 0, 4)
normal_run(p4, 'Capstone with Lityx (AI/ML BI provider): applied ML models to a manufacturing process, reducing predicted waste by 40% across 5M+ rows of simulated data.', size=10)

add_hr(doc)

# ============================================================
# ADDITIONAL
# ============================================================
section_heading(doc, 'Additional')

p = doc.add_paragraph()
tight(p, 2, 0)
bold_run(p, 'President', size=10)
normal_run(p, '  |  Colorado Sled Hockey  |  Sep 2023 – Present', size=10)

p2 = doc.add_paragraph()
tight(p2, 0, 0)
normal_run(p2, '501(c)(3) nonprofit providing hockey opportunities for children and adults with physical disabilities. Leads club activities, tournaments, fundraising, and player development.', size=10)

# ============================================================
# SAVE
# ============================================================
out_path = r'C:\Users\JPHeb\Documents\claude-projects\job-search\JP_Hebert_Resume_V3.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
